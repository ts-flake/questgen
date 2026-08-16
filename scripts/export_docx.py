r"""M5: export selected bank entries -> worksheet .docx

Layout: a grid table spanning the text width, 3 columns when marks are shown
(number | content | marks) and 2 otherwise. Each sub-part is ONE row whose cell holds the
text and then the answer-writing lines, so nothing blank sits between sub-parts; the mark
sits one line above the bottom of its row, and a final row merges the last two columns for
"[Total: X]". A teacher copy drops the writing space and puts each answer in its own red
row directly under the row it answers — question line, answer line.

The writing space comes from `answer_area` (schema v2), so the blank keeps whatever the
paper printed beside it ("[ANSWER] km", "$[ANSWER]", "equation: [ANSWER] / conditions:
[ANSWER]"); a blank still sitting inline in the text is rendered by the text, so v1 rows
export unchanged. Answers likewise live on the sub-parts once a question has any, which is
why every answer path goes through `answer_lines`.

Math: entries keep MinerU's latex ($ / $$). Word cannot render that, so latex is converted
to real OMML equations (latex2mathml -> mathml2omml); a unicode transliteration is only the
fallback when conversion fails. Chemistry \ce{} is expanded before conversion. HTML tables
become real docx tables, with rowspan/colspan honoured via cell merges.
"""
from __future__ import annotations

import re
import shutil
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from docx.shared import Cm, Emu, Pt, RGBColor

ANS_RED = RGBColor(0xC0, 0x00, 0x00)      # teacher-copy inline answers/solutions


# w:rPr children that must precede w:color (CT_RPr is an ordered sequence). Anything not
# listed sorts after, so inserting before the first unlisted child keeps the order valid.
_RPR_BEFORE_COLOR = tuple(qn("w:" + t) for t in (
    "rStyle", "rFonts", "b", "bCs", "i", "iCs", "caps", "smallCaps", "strike", "dstrike",
    "outline", "shadow", "emboss", "imprint", "noProof", "snapToGrid", "vanish", "webHidden"))


def _set_red(cell):
    """Colour every run in a cell red (teacher copy) — the wording AND the equations.

    An OMML equation's runs are `m:r`, not `w:r`, so `paragraph.runs` never sees them and a
    latex answer exported as a native equation stayed black next to its red text. Their colour
    lives in a `w:rPr` inside the `m:r`, which sits after `m:rPr` and before the `m:t`."""
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.color.rgb = ANS_RED
    for mr in cell._tc.findall(".//" + qn("m:r")):
        rpr = mr.find(qn("w:rPr"))
        if rpr is None:
            rpr = mr.makeelement(qn("w:rPr"), {})
            mrpr = mr.find(qn("m:rPr"))
            mrpr.addnext(rpr) if mrpr is not None else mr.insert(0, rpr)
        color = rpr.find(qn("w:color"))
        if color is None:
            color = rpr.makeelement(qn("w:color"), {})
            after = [c for c in rpr if c.tag not in _RPR_BEFORE_COLOR]
            after[0].addprevious(color) if after else rpr.append(color)
        color.set(qn("w:val"), str(ANS_RED))

import context

try:  # latex → native Word equations; unicode downconversion is the fallback
    import latex2mathml.converter as _l2m
    import mathml2omml as _m2o
    HAS_OMML = True
except Exception:
    HAS_OMML = False

OMML_NS = 'xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"'


# ---------------------------------------------------------------- mhchem \ce{}
# MinerU never emits \ce{}; the chem clean pass adds it. latex2mathml doesn't know mhchem,
# so we expand \ce{...} to plain LaTeX (upright element symbols, subscripts, charge
# superscripts, reaction arrows) before conversion. Covers school-chemistry notation.
def _ce_body_to_latex(s: str) -> str:
    s = s.replace("<=>", r" \rightleftharpoons ").replace("<->", r" \leftrightarrow ")
    s = s.replace("->", r" \rightarrow ").replace("<-", r" \leftarrow ")
    out, i, n = [], 0, len(s)
    while i < n:
        c = s[i]
        if c == "\\":                                    # a LaTeX command (e.g. our arrows) — keep intact
            j = i + 1
            while j < n and s[j].isalpha():
                j += 1
            out.append(s[i:j] if j > i + 1 else s[i:i + 2]); i = max(j, i + 2)
        elif c.isalpha():                                # element symbols -> upright
            j = i
            while j < n and s[j].isalpha():
                j += 1
            out.append(r"\mathrm{" + s[i:j] + "}"); i = j
        elif c.isdigit():                                # digit after a group/) = subscript, else coefficient
            j = i
            while j < n and s[j].isdigit():
                j += 1
            prev = out[-1] if out else ""
            out.append(("_{" + s[i:j] + "}") if prev.endswith(("}", ")")) else s[i:j]); i = j
        elif c == "^":                                   # charge / superscript: ^{..} or ^2- / ^+
            i += 1
            if i < n and s[i] == "{":
                k = s.find("}", i); k = n if k < 0 else k
                out.append("^{" + s[i + 1:k] + "}"); i = k + 1
            else:
                j = i
                while j < n and s[j].isdigit():
                    j += 1
                if j < n and s[j] in "+-":
                    j += 1
                out.append("^{" + s[i:j] + "}"); i = j
        elif c == "_":                                   # existing subscript _n or _{..}
            i += 1
            if i < n and s[i] == "{":
                k = s.find("}", i); k = n if k < 0 else k
                out.append("_{" + s[i + 1:k] + "}"); i = k + 1
            else:
                j = i
                while j < n and s[j].isalnum():
                    j += 1
                out.append("_{" + s[i:j] + "}"); i = j
        else:
            out.append(c); i += 1
    return "".join(out)


_CE = re.compile(r"\\ce\s*\{")


def expand_ce(s: str) -> str:
    if not s or "\\ce" not in s:
        return s
    out, i = [], 0
    while True:
        m = _CE.search(s, i)
        if not m:
            out.append(s[i:]); break
        out.append(s[i:m.start()])
        depth, j = 1, m.end()
        while j < len(s) and depth:
            depth += 1 if s[j] == "{" else -1 if s[j] == "}" else 0
            j += 1
        out.append(_ce_body_to_latex(s[m.end():j - 1])); i = j
    return "".join(out)


GRID_ENV = re.compile(r"\\begin\{(array|[bBpvV]?matrix|smallmatrix|cases)\}")


def omml_el(latex: str):
    latex = expand_ce(latex)                             # mhchem \ce{} -> plain LaTeX
    # aligned/align environments carry alignment '&' which latex2mathml emits as a bare
    # '<mi>&</mi>' — invalid XML that breaks the mathml→omml SAX parse. Drop those tabs.
    # BUT array/matrix/cases use '&' as real COLUMN separators (latex2mathml -> <mtable>),
    # so keep '&' there — stripping it collapses an N-column grid to one column.
    if not GRID_ENV.search(latex):
        latex = re.sub(r"(?<!\\)&", "", latex)
    o = _m2o.convert(_l2m.convert(latex))
    return parse_xml(o.replace("<m:oMath>", f"<m:oMath {OMML_NS}>", 1))

IMG_MARK = re.compile(r"!\[\]\(([^)]+)\)")
TABLE_HTML = re.compile(r"<table>.*?</table>", re.S)


def iter_part_nodes(parts, level=0, prefix=""):
    """Nested parts tree -> (node, level, path) in reading order. `path` is the composite
    label ('(b)(i)') — the same key `answer_lines` reports, so the two can be matched up."""
    for p in parts or []:
        path = prefix + (p.get("no") or "")
        yield p, level, path
        yield from iter_part_nodes(p.get("children", []), level + 1, path)


_LETTERS = "ABCDEFGHIJ"


def option_index(key) -> int:
    """0-based position of an option key. Internally keys are '(1)'..'(n)'; legacy
    A/B/C/D or bare digits are still accepted."""
    k = str(key).strip().strip("()（）").strip().rstrip(".、")
    if k.isdigit():
        return max(0, int(k) - 1)
    return _LETTERS.index(k.upper()) if k[:1].upper() in _LETTERS else 0


def option_label(key: str, fmt: str) -> str:
    """Re-render an option key in the chosen style. fmt = '<letter|num>_<paren|dot|bare>'.
    The stored key is always '(n)' — the display format is an export choice."""
    idx = option_index(key)
    lab = _LETTERS[idx] if fmt.startswith("letter") and idx < len(_LETTERS) else str(idx + 1)
    if fmt.endswith("paren"):
        return f"({lab})"
    if fmt.endswith("dot"):
        return f"{lab}."
    return lab


BLANK_FILL = "_" * 12          # set per export by build_docx (underscore vs dot-leader)
# figure/table caption options, set per export by build_docx (see CAPTION_DEFAULTS)
CAPTION: dict = {}
CAPTION_DEFAULTS = {
    "style": "bold",        # "italic" | "bold" | "plain" | "none" (none = no captions)
    "fig_word": "Fig.",     # word used for figures ("Fig." / "Figure" / "Diagram" / "图")
    "table_word": "Table",  # word used for tables
    "tables": True,         # caption tables too, or figures only
    "number": "per_question",   # "per_question" (1.2) | "continuous" (7) | "none"
}


def subst_placeholders(s: str, qno) -> str:
    """[QN] -> the question's displayed number. [ANSWER] is left in place; it is
    rendered (right-aligned blank) at write time by write_rich_line."""
    return s.replace("[QN]", str(qno)) if s else s


# ---------------------------------------------------------------- latex → text

LATEX_SUB = [
    (r"\\times", "×"), (r"\\div", "÷"), (r"\\angle", "∠"), (r"\\circ", "°"),
    (r"\\%", "%"), (r"\\$", "$"), (r"\\quad", " "), (r"\\;|\\,|\\!", " "),
    (r"\\mathrm\{([^}]*)\}", r"\1"), (r"\\text\{([^}]*)\}", r"\1"),
    # emphasis (toolbar B/I/U): OMML renders these bold/italic; this is the no-OMML fallback
    # — keep the text, drop the command name (bold/italic can't be shown in plain unicode)
    (r"\\text(?:bf|it|rm|sf|tt)\{([^}]*)\}", r"\1"), (r"\\math(?:bf|it|rm|sf|bb|cal)\{([^}]*)\}", r"\1"),
    (r"\\emph\{([^}]*)\}", r"\1"),
    (r"\\overline\{([^}]*)\}", r"\1"), (r"\\underline\{([^}]*)\}", r"_\1_"),
    (r"\\sqrt\{([^}]*)\}", r"√(\1)"), (r"\\pi", "π"), (r"\\approx", "≈"),
    (r"\\rightleftharpoons", "⇌"), (r"\\leftrightarrow", "↔"),
    (r"\\rightarrow|\\to", "→"), (r"\\leftarrow", "←"),
    (r"\\neq", "≠"), (r"\\leq|\\le", "≤"), (r"\\geq|\\ge", "≥"),
    (r"\\begin\{array\}\{[^}]*\}", ""), (r"\\end\{array\}", ""),
    (r"\\begin\{(?:aligned|align|split|gathered|gather)\*?\}", ""),
    (r"\\end\{(?:aligned|align|split|gathered|gather)\*?\}", ""),
    (r"\\\\", "\n"), (r"&", ""), (r"\\left|\\right", ""),
    (r"~", " "), (r"\\cents?\b", "¢"), (r"\\heartsuit", "♥"), (r"\\cdot", "·"), (r"\\pm", "±"), (r"\\star", "★"),
    (r"\\bullet", "•"), (r"\\degree", "°"), (r"\\triangle", "△"), (r"\\square", "□"),
]


def latex_to_text(s: str) -> str:
    s = expand_ce(s)                         # mhchem \ce{} -> plain LaTeX (unicode fallback path)
    def frac(m):
        return f"{m.group(1)}/{m.group(2)}"
    prev = None
    while prev != s:  # nested fracs
        prev = s
        s = re.sub(r"\\[dt]?frac\{([^{}]*)\}\{([^{}]*)\}", frac, s)
    for pat, rep in LATEX_SUB:
        s = re.sub(pat, rep, s)
    s = re.sub(r"\^\{([^}]*)\}", r"^\1", s)
    s = re.sub(r"_\{([^}]*)\}", r"\1", s)
    s = s.replace("{", "").replace("}", "")
    return s


def clean_text(s: str) -> str:
    """Latex inside $..$/$$..$$ → unicode; keep the rest verbatim.
    Escaped currency dollars (\\$) must not participate in math-span pairing."""
    s = s.replace("\\$", "\x00")
    s = re.sub(r"\$\$([^$]*)\$\$", lambda m: latex_to_text(m.group(1)).strip(), s)
    s = re.sub(r"\$([^$]*)\$", lambda m: latex_to_text(m.group(1)).strip(), s)
    s = s.replace("\x00", "$").replace("\\_", "_")
    return re.sub(r"\n{3,}", "\n\n", s).strip()


# $..$ / $$..$$ math, plus bare \begin{env}..\end{env} blocks (aligned, cases, …)
# that segmenters sometimes emit without $ delimiters.
MATH_SPAN = re.compile(
    r"(\$\$[^$]*\$\$|\$[^$]*\$|\\begin\{[a-zA-Z*]+\}.*?\\end\{[a-zA-Z*]+\})", re.S)


# inline equations must not butt against adjacent words. We add a thin space run
# before/after an equation when a real word (not whitespace/closing punctuation) sits
# next to it — so "$v$at" renders "v at", never "vat".
PUNCT_AFTER = set(".,;:!?)]}%’\"")


def _para_has_content(p) -> bool:
    return bool(p.runs) or bool(p._p.findall(qn("m:oMath")))


def _space_before_math(p) -> bool:
    """A leading space is needed if the paragraph already has content that does not
    already end in whitespace."""
    if not _para_has_content(p):
        return False
    return not (p.runs and (p.runs[-1].text or "").endswith((" ", "\t")))


def _space_after_math(next_tok: str, stripped_next: bool) -> bool:
    """Whether to add a trailing space after an equation. `stripped_next` = the caller
    will strip leading whitespace off the following text (so we must supply the gap)."""
    s = next_tok.replace("\x00", "$")
    if stripped_next:
        s = s.lstrip(" \t")
    if not s or s[0] == "\n":            # new paragraph next → no trailing space
        return False
    if not stripped_next and s[0].isspace():
        return False                     # a real space already separates them
    return s[0] not in PUNCT_AFTER


# An answer line the blank OPENS ("[ANSWER] km", "$[ANSWER]") is right-aligned so the rule
# ends at the right margin. One the blank only ends ("equation: [ANSWER]", or a sentence with
# a blank in it) stays left, or right-aligning would drag its label off to the margin too.
ANSWER_LEAD = re.compile(r"^\s*\\?[$¢£€]?\s*\[ANSWER\]")


def write_rich_line(para, line: str):
    """One line of text into a paragraph: $..$ spans (and bare latex environments)
    become native equations (OMML) when converters are available, else unicode
    fallback. A line the [ANSWER] blank opens is right-aligned (exam answer line)."""
    if "[ANSWER]" in line:
        if ANSWER_LEAD.match(line):
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        line = line.replace("[ANSWER]", BLANK_FILL)
    line = line.replace("\\$", "\x00")
    toks = [t for t in MATH_SPAN.split(line) if t]
    for ti, tok in enumerate(toks):
        if tok.startswith("$") or tok.startswith("\\begin{"):
            latex = (tok.strip("$") if tok.startswith("$") else tok).replace("\x00", "\\$").strip()
            _put_math(para, latex, lead_space=_space_before_math(para))
            nxt = toks[ti + 1] if ti + 1 < len(toks) else ""
            if _space_after_math(nxt, stripped_next=False):
                para.add_run(" ")
        else:
            para.add_run(tok.replace("\x00", "$").replace("\\_", "_"))


def _cell_para(cell, first: bool):
    """Reuse the cell's empty first paragraph, else append a fresh one."""
    p0 = cell.paragraphs[0]
    if first and not p0.text and not p0.runs and not p0._p.findall(qn("m:oMath")):
        return p0
    return cell.add_paragraph()


# aligned/align/split/gather rows are stacked, one display equation per row. OMML does
# not honour the environment's line breaks, so we split on \\ and lay out a paragraph each.
ALIGN_ENV = re.compile(
    r"\\begin\{(aligned|align|alignat|split|gathered|gather)\*?\}(.*?)\\end\{\1\*?\}", re.S)


def _math_rows(tok: str) -> list[str]:
    """A math token -> list of per-line latex fragments (aligned envs split on \\\\)."""
    latex = tok.strip("$").strip() if tok.startswith("$") else tok.strip()
    m = ALIGN_ENV.search(latex)
    if not m:
        return [latex]
    rows = [r.strip() for r in re.split(r"\\\\", m.group(2)) if r.strip()]
    return rows or [latex]


def _put_math(para, latex: str, lead_space: bool = False):
    """Append one equation (native OMML, else unicode) to a paragraph, optionally
    preceded by a space so it doesn't stick to the word before it."""
    if lead_space:
        para.add_run(" ")
    if HAS_OMML and latex:
        try:
            para._p.append(omml_el(latex))
            return
        except Exception:
            pass
    para.add_run(latex_to_text(latex))


def write_rich_text(cell, text: str, first: bool):
    """Multi-line text into a cell as real paragraphs (Word ignores \\n). Math blocks
    ($$..$$, $..$, or bare \\begin{env}..\\end{env}) may themselves span newlines, so
    they are kept intact and only the surrounding plain text is split on newlines.
    Aligned environments are laid out one equation per line."""
    para = None
    # protect escaped '\$' (literal dollar) so it can't pair with a real $ math delimiter
    text = (text or "").replace("\\$", "\x00")
    toks = [t for t in MATH_SPAN.split(text) if t]
    for ti, tok in enumerate(toks):
        if tok.startswith("$") or tok.startswith("\\begin{"):
            for j, row in enumerate(_math_rows(tok.replace("\x00", "\\$"))):
                if j > 0:
                    para = None                 # each aligned row on its own line
                if para is None:
                    para = _cell_para(cell, first); first = False
                _put_math(para, row, lead_space=(j == 0 and _space_before_math(para)))
            # plain lines are stripped below, so supply the trailing gap ourselves
            nxt = toks[ti + 1] if ti + 1 < len(toks) else ""
            if _space_after_math(nxt, stripped_next=True):
                para.add_run(" ")
        else:
            lines = tok.split("\n")
            for i, ln in enumerate(lines):
                if i > 0:
                    para = None                 # a newline forces the next paragraph
                ln = ln.strip()
                if para is None:
                    para = _cell_para(cell, first); first = False
                if ln:
                    write_rich_line(para, ln.replace("\x00", "\\$"))   # restore; write_rich_line re-protects
                # an empty ln leaves `para` as an empty paragraph — a preserved blank line


# ---------------------------------------------------------------- cell writers

def _cell_text(v: str) -> str:
    """Inner HTML of a <td> -> text with math kept ($...$): <br> -> newline, other tags dropped."""
    v = re.sub(r"(?i)<br\s*/?>", "\n", v)
    v = re.sub(r"<[^>]+>", "", v)
    return v.strip()


def _parse_html_table(html: str):
    """[[cell dicts]] with colspan/rowspan per row (th treated like td)."""
    out = []
    for rh in re.findall(r"(?is)<tr[^>]*>(.*?)</tr>", html):
        cells = []
        for m in re.finditer(r"(?is)<t[dh]([^>]*)>(.*?)</t[dh]>", rh):
            attrs, content = m.group(1), m.group(2)
            sp = lambda a: int((re.search(rf'{a}\s*=\s*"?(\d+)', attrs, re.I) or [0, "1"])[1])
            cells.append({"content": content, "colspan": max(1, sp("colspan")),
                          "rowspan": max(1, sp("rowspan"))})
        if cells:
            out.append(cells)
    return out


def add_html_table(cell, html: str):
    """HTML table -> a real docx table: rowspan/colspan honoured via cell merges, and each
    cell's LaTeX ($...$) rendered as math (not crude unicode)."""
    rows = _parse_html_table(html)
    if not rows:
        return
    # place cells on a grid, tracking occupancy so spans push later cells right/down
    occ, placed, ncols = set(), [], 0
    for r, cells in enumerate(rows):
        c = 0
        for cd in cells:
            while (r, c) in occ:
                c += 1
            rs, cs = cd["rowspan"], cd["colspan"]
            placed.append((r, c, rs, cs, cd["content"]))
            for rr in range(r, r + rs):
                for cc in range(c, c + cs):
                    occ.add((rr, cc))
            c += cs
        ncols = max(ncols, c)
    nrows = len(rows)
    if not ncols:
        return
    t = cell.add_table(rows=nrows, cols=ncols)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r, c, rs, cs, content in placed:
        r2, c2 = min(r + rs - 1, nrows - 1), min(c + cs - 1, ncols - 1)
        anchor = t.cell(r, c)
        if r2 > r or c2 > c:
            anchor = anchor.merge(t.cell(r2, c2))       # rowspan/colspan
        anchor.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        write_rich_text(anchor, _cell_text(content), first=True)   # $...$ -> real math


def _caption(cell, kind: str, figctx) -> None:
    """Centred caption for a figure/table, with ONE blank line above and below it. Called
    AFTER a figure (caption sits below it) and BEFORE a table (title sits above it).
    Per the export's CAPTION options; counters advance even when captions are off so
    numbering stays stable."""
    if not figctx:
        return
    figctx[kind] = figctx.get(kind, 0) + 1
    cfg = CAPTION or CAPTION_DEFAULTS
    style = cfg.get("style", "italic")
    if style == "none" or (kind == "table" and not cfg.get("tables", True)):
        return
    word = cfg.get("table_word", "Table") if kind == "table" else cfg.get("fig_word", "Fig.")
    mode = cfg.get("number", "per_question")
    if mode == "none":
        num = ""
    elif mode == "continuous":
        run = figctx.setdefault("_run", {})         # shared across questions (whole doc)
        run[kind] = run.get(kind, 0) + 1
        num = str(run[kind])
    else:
        num = f"{figctx['qno']}.{figctx[kind]}"
    cell.add_paragraph()                             # blank line above the caption
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"{word} {num}".strip())
    if style == "bold":
        r.bold = True
    elif style == "italic":
        r.italic = True                              # font size inherits Normal (same as body)
    cell.add_paragraph()                             # blank line below the caption


def add_content(cell, text: str, img_dirs: list[Path], imgs_meta: list, figctx=None):
    """Write text into cell, replacing image markers / html tables inline. Figures and
    tables are centred and get a 'Fig./Table QN.k' caption (when figctx given)."""
    text = text or ""
    pat = re.compile(r"(!\[\]\([^)]+\)|<table>.*?</table>)", re.S)
    first = True
    for seg in pat.split(text):
        seg = seg.strip()
        if not seg:
            continue
        m = IMG_MARK.fullmatch(seg)
        if m:
            p = next((d / m.group(1) for d in img_dirs if (d / m.group(1)).is_file()), None)
            if p:
                p0 = cell.paragraphs[0]
                reuse = first and not p0.text and not p0.runs and not p0._p.findall(qn("m:oMath"))
                para = p0 if reuse else cell.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER   # figures centred
                run = para.add_run()
                try:
                    from PIL import Image
                    w, h = Image.open(p).size
                    wcm, hcm = w / 60, h / 60          # natural size @ ~150dpi
                    scale = min(1.0, 9.0 / wcm, 3.7 / hcm)  # max 9cm wide, 3.7cm tall
                    width = Cm(max(2.0, wcm * scale))
                except Exception:
                    width = Cm(7)
                run.add_picture(str(p), width=width)
                _caption(cell, "fig", figctx)
            first = False
        elif seg.startswith("<table>"):
            _caption(cell, "table", figctx)          # table title goes ABOVE the table
            add_html_table(cell, seg)
            first = False
        else:
            write_rich_text(cell, seg, first)
            first = False


# ---------------------------------------------------------------- document

def _grid(doc, widths):
    t = doc.add_table(rows=0, cols=len(widths))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    from docx.oxml.ns import qn
    tblPr = t._tbl.tblPr
    tblPr.append(tblPr.makeelement(qn("w:tblLayout"), {qn("w:type"): "fixed"}))
    t._widths = widths
    return t


def _row(t, no: str = "", marks=None, min_height=None):
    row = t.add_row()
    cells = row.cells
    for c, w in zip(cells, t._widths):
        c.width = w
        tcW = c._tc.tcPr.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcW")
        if tcW is not None:
            tcW.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type", "dxa")
    if min_height is not None:                       # guaranteed answer space (survives Word layout)
        row.height = min_height
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    if no:
        cells[0].paragraphs[0].add_run(no).bold = True      # question number in bold
    if len(cells) >= 3 and marks is not None:
        cells[2].vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
        p = cells[2].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(f"[{marks}]")
        blank = cells[2].add_paragraph()                 # [N] + trailing blank, bottom-aligned, so
        blank.alignment = WD_ALIGN_PARAGRAPH.CENTER      # the mark sits one line above the bottom
    return cells[1]


def add_part(cell, no: str, text: str, lvl: int, dirs, sub, figctx=None):
    """Render a sub-part as a hanging-indent 'list' item: the label hangs at the
    level's indent, wrapped text aligns under the text (Word list layout)."""
    indent = Cm(0.75 * (lvl + 1))
    label_written = False
    for seg in re.split(r"(!\[\]\([^)]+\)|<table>.*?</table>)", text or "", flags=re.S):
        seg = seg.strip()
        if not seg:
            continue
        if IMG_MARK.fullmatch(seg) or seg.startswith("<table>"):
            add_content(cell, seg, dirs, [], figctx)
            continue
        for ln in seg.split("\n"):
            ln = ln.strip()
            if not ln:
                continue
            p0 = cell.paragraphs[0]
            para = p0 if (not label_written and not p0.text and not p0.runs) else cell.add_paragraph()
            para.paragraph_format.left_indent = indent
            if not label_written:
                para.paragraph_format.first_line_indent = Cm(-0.75)   # hanging label
                para.add_run(no + "  ").bold = True                   # sub-part label in bold
                label_written = True
            write_rich_line(para, sub(ln))
    if not label_written and no:            # empty-text part (synthesized "(a)" parent): label only
        p0 = cell.paragraphs[0]
        para = p0 if (not p0.text and not p0.runs) else cell.add_paragraph()
        para.paragraph_format.left_indent = indent
        para.paragraph_format.first_line_indent = Cm(-0.75)
        para.add_run(no).bold = True


def entry_img_dirs(ctx: context.Ctx, entry: dict) -> list[Path]:
    # cross-source export: the caller attaches _img_dirs computed from THIS entry's own
    # source, so images resolve correctly even when entries come from several sources.
    if entry.get("_img_dirs"):
        return [Path(d) for d in entry["_img_dirs"]]
    stem = Path(entry["meta"]["file"]).stem
    # solution images live in the separate <stem>_ans extraction
    return [ctx.extracted_dir / stem, ctx.extracted_dir / f"{stem}_ans"]


def answer_display(entry: dict, val: str, fmt: str) -> str:
    """An MCQ answer is stored as its option label ('(3)'). Re-render it in the export's
    chosen option style so the printed answer matches the printed options. Anything that
    is not a pure label reference — a computed value, an expression, per-part answers — is
    returned untouched, including a bare number that is one of the option VALUES."""
    opts = entry.get("options") or {}
    if not opts:
        return val
    toks = [t for t in re.split(r"[\s,;/、]+", val.strip()) if t]
    if not toks:
        return val
    keys = {str(k).strip(): k for k in opts}
    bare = {str(k).strip().strip("()（）"): k for k in opts}
    values = {str(v).strip() for v in opts.values()}
    hits = []
    for t in toks:
        if t in keys:                                  # exact label, e.g. "(3)"
            hits.append(keys[t])
            continue
        b = t.strip("()（）")
        if b in bare and t not in values:              # bare index that isn't a value
            hits.append(bare[b])
            continue
        return val
    return ", ".join(option_label(k, fmt) for k in hits)


def add_options(cell, options: dict, fmt: str, dirs, sub):
    """One paragraph per MCQ option: bold label (chosen style) + value (latex/img)."""
    for k in sorted(options, key=option_index):
        para = cell.add_paragraph()
        run = para.add_run(option_label(k, fmt) + " ")
        run.bold = True
        val = sub(str(options[k]))
        if IMG_MARK.search(val) or val.strip().startswith("<table>"):
            add_content(cell, val, dirs, [])     # figure/table option -> own block
        else:
            write_rich_line(para, val)


# ---------------------------------------------------------------- sections & marks
# Section grouping = one exam-style block per question type (MCQ first), sorted by marks
# within. Numbering stays continuous (1..N across the whole paper) so [QN]/figure refs and
# the answers section line up with the question numbers.
TYPE_ORDER = ["mcq", "true_false", "fill_blank", "short_answer", "word_problem", "structured"]
TYPE_NAME = {"mcq": "Multiple Choice", "true_false": "True / False",
             "fill_blank": "Fill in the Blanks", "short_answer": "Short Answer",
             "word_problem": "Word Problems", "structured": "Structured Questions"}


def entry_type(e: dict) -> str:
    """Problem-type used for section grouping: the tag if present, else inferred from shape."""
    t = (e.get("tags") or {}).get("type")
    if t:
        return t
    return "mcq" if e.get("options") else ("structured" if e.get("parts") else "short_answer")


def entry_marks(e: dict):
    m = e.get("meta", {}).get("marks")
    return m if isinstance(m, (int, float)) else None


def total_marks(entries: list[dict]) -> int:
    return sum(m for e in entries if (m := entry_marks(e)) is not None)


def _marks_key(e: dict):
    m = entry_marks(e)
    return (m is None, m or 0)                       # marked ascending, then unmarked (stable)


def order_sections(entries: list[dict]) -> list[tuple[str, str, list[dict]]]:
    """Group entries by type, order the groups (TYPE_ORDER, then any extra types in first-seen
    order), sort each group by marks ascending (stable → cart order within equal marks).
    Returns [(section_label 'A'/'B'/…, display name, [entries]), …]."""
    groups: dict = {}
    seen: list = []
    for e in entries:
        ty = entry_type(e)
        if ty not in groups:
            groups[ty] = []
            seen.append(ty)
        groups[ty].append(e)
    order = sorted(seen, key=lambda ty: TYPE_ORDER.index(ty) if ty in TYPE_ORDER
                   else len(TYPE_ORDER) + seen.index(ty))
    out = []
    for k, ty in enumerate(order):
        name = TYPE_NAME.get(ty, ty.replace("_", " ").title())
        out.append((chr(65 + k), name, sorted(groups[ty], key=_marks_key)))
    return out


ANS_LINES = 2                                            # answer-writing lines below each question


def _answer_lines(cell, text: str, area: str | None = None, sub=None, indent=None) -> None:
    """Gold layout: the answer-writing space lives INLINE in the question/part cell (no separate
    answer row). Every question gets ANS_LINES writing lines; a rendered [ANSWER] placeholder is
    already one such (right-aligned underscore) line, so it counts toward the total and the
    padding blanks match its right alignment. Otherwise the blanks are left at default.

    Schema v2 lifts the TRAILING blank out of the text into `answer_area`, keeping whatever the
    paper prints beside it — a unit ("[ANSWER] km"), a symbol ("$[ANSWER]") or several labelled
    blanks ("equation: [ANSWER]\\n\\nconditions: [ANSWER]"). That template is printed here, one
    paragraph per line, so the exported blank still carries its unit. A blank left mid-sentence
    in the text is not part of the area and is rendered by the text itself, as before."""
    written, right = 0, False
    for ln in ((area or "").split("\n") if (area or "").strip() else []):
        p = cell.add_paragraph()
        if indent is not None:
            p.paragraph_format.left_indent = indent
        if ln.strip():
            write_rich_line(p, sub(ln) if sub else ln)   # aligns itself, see ANSWER_LEAD
            right = right or p.alignment == WD_ALIGN_PARAGRAPH.RIGHT
            written += 1
    right = right or any(ANSWER_LEAD.match(ln) for ln in (text or "").split("\n"))
    used = written + (text or "").count("[ANSWER]")
    for _ in range(max(0, ANS_LINES - used)):
        p = cell.add_paragraph()
        if indent is not None:
            p.paragraph_format.left_indent = indent
        if right:                                        # padding matches the rule it extends
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def _question_total(e: dict):
    m = entry_marks(e)
    if m is not None:
        return m
    s = sum((n.get("marks") or 0) for n, _, _ in iter_part_nodes(e.get("parts", [])))
    return s or None


def _total_row(t, e: dict) -> None:
    """Dedicated bottom row: merge content+marks cols and right-align a bold '[Total: X]'
    with a blank line above and below it (too long for the 1.5cm marks col; merged it has
    room, right edge still lines up)."""
    qtotal = _question_total(e)
    if qtotal is None:                                    # no marks to total: no row at all,
        return                                            # an empty one just read as a stray gap
    _row(t)
    cells = t.rows[-1].cells
    merged = cells[1].merge(cells[2])
    above = merged.paragraphs[0]                          # blank line above
    above.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    mid = merged.add_paragraph()
    mid.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    mid.add_run(f"[Total: {qtotal}]").bold = True
    below = merged.add_paragraph()                        # blank line below
    below.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def answer_lines(e: dict) -> list[tuple[str, str, str]]:
    """[(label, answer, solution)] for one entry, flattened in reading order.

    Answers live on the parts once a question has sub-parts, so an exporter cannot read
    `entry["answer"]` alone: it holds a value only for a question with no parts (or one whose
    key named no part at all). Label is "" for the entry-level row."""
    out = []
    ans = (e.get("answer") or {}).get("value") or ""
    if ans or e.get("solution"):
        out.append(("", str(ans), str(e.get("solution") or "")))

    def walk(parts, pre=""):
        for p in parts or []:
            label = pre + (p.get("no") or "")
            if p.get("answer") or p.get("solution"):
                out.append((label, str(p.get("answer") or ""), str(p.get("solution") or "")))
            walk(p.get("children") or [], label)

    walk(e.get("parts"))
    return out


def _question_block(t, i: int, e: dict, ctx, mcq_label: str, marks_col: bool, running: dict,
                    teacher: bool = False):
    """Render one question with ONE canonical skeleton (aligned to the gold sample.docx), so
    every export mode lines up:

        stem row  ->  for each sub-part: ONE row = [ label+text + inline answer lines | mark ]
                  ->  [Total: X] row

    The answer-writing space is inline in each sub-part's own cell (no separate answer row, so
    no blank row between sub-parts). Modes vary only in:
      * marks_col: sub-part mark shown in col3 + the [Total: X] row emitted; without it, no
        marks/total column at all.
      * teacher:   the inline answer space is replaced by the red answer/solution block."""
    dirs = entry_img_dirs(ctx, e)
    sub = lambda s: subst_placeholders(s, i)
    figctx = {"qno": i, "fig": 0, "table": 0, "_run": running}
    cell = _row(t, f"{i}.")                           # question total -> [Total: X] row, not the col
    add_content(cell, sub(e["stem"]), dirs, e.get("imgs", []), figctx)
    if e.get("options"):
        add_options(cell, e["options"], mcq_label, dirs, sub)
    parts = list(iter_part_nodes(e.get("parts", [])))
    if parts:
        cell.add_paragraph()                          # blank line below the stem, before the sub-parts

    if teacher:                                       # marking copy: red answer row under its question
        slots = {lab: (ans, sol) for lab, ans, sol in answer_slots(e)}
        if "" in slots:                               # a question that asks in its own right
            _answer_pair(t, e, i, mcq_label, dirs, *slots[""])          # pair under the stem
        for node, lvl, path in parts:
            cell = _row(t, "", marks=node.get("marks") if marks_col else None)
            add_part(cell, node.get("no", ""), node.get("text", ""), lvl, dirs, sub, figctx)
            cell.add_paragraph()                      # trailing line so the mark sits one above bottom
            if path in slots:
                _answer_pair(t, e, i, mcq_label, dirs, *slots[path],
                             indent=Cm(0.75 * (lvl + 1)))
        if marks_col:
            _total_row(t, e)                          # inline-answer copy also gets the total
        return

    # student copy: each sub-part is one row with its answer space inline in the same cell.
    if not parts:
        _answer_lines(cell, e.get("stem", ""), e.get("answer_area"), sub)
    for node, lvl, path in parts:
        cell = _row(t, "", marks=node.get("marks") if marks_col else None)
        add_part(cell, node.get("no", ""), node.get("text", ""), lvl, dirs, sub, figctx)
        _answer_lines(cell, node.get("text", ""), node.get("answer_area"), sub,
                      indent=Cm(0.75 * (lvl + 1)))
    if marks_col:
        _total_row(t, e)                              # dedicated [Total: X] row


MISSING = "null"                    # printed where an answer slot is empty, see _answer_pair


def answer_slots(e: dict) -> list[tuple[str, str, str]]:
    """[(label, answer, solution)] for every slot that ASKS something, in reading order.

    An answer slot belongs to the question itself when it has no sub-parts, otherwise to each
    LEAF sub-part — a parent part is only an intro line, so it owns a slot solely when the key
    really put an answer on it. Slots with nothing recorded are still returned, with empty
    strings, because both output paths print a fixed pair and show the gap (see _answer_pair).
    `answer_lines` stays the definition of where an answer lives; this adds where one is due."""
    amap = {lab: (a, s) for lab, a, s in answer_lines(e)}
    parts = list(iter_part_nodes(e.get("parts", [])))
    out = []
    if not parts or "" in amap:
        out.append(("", *amap.get("", ("", ""))))
    for node, _lvl, path in parts:
        if not node.get("children") or path in amap:
            out.append((path, *amap.get(path, ("", ""))))
    return out


def _prefix_bold(cell, label: str) -> None:
    """Put a bold label at the very start of a cell's first paragraph, ahead of whatever
    add_content already wrote there (a text run, an equation or an image)."""
    p = cell.paragraphs[0]
    run = p.add_run(label)
    run.bold = True
    pPr = p._p.find(qn("w:pPr"))
    pPr.addnext(run._r) if pPr is not None else p._p.insert(0, run._r)


def _answer_pair(t, e: dict, i: int, mcq_label: str, dirs, ans: str = "", sol: str = "",
                 indent=None, label: str = "", no: str = "") -> None:
    """One answer slot as its fixed two red rows: bold 'Ans:' then bold 'Solution:', with
    MISSING printed for an empty one.

    Collapsing the empty ones is what made the marking copy ragged — a question could show one
    row, two, or two plus a gap — so a marker could not tell "no solution recorded" from a
    different row layout. A visible `null` says which slots still need filling in.

    Both output paths share this, so they cannot drift: the inline marking copy passes
    `indent` to line the pair up under its own sub-part and no `label`, because the question
    text is on the row above; the answers table at the end passes `label` (the sub-part path)
    and `no` (the question number, on the first row only), because neither is next to it."""
    sub = lambda s: subst_placeholders(s, i)
    pre = f"{label} " if label else ""

    def _close(c):
        if indent is not None:
            for p in c.paragraphs:
                p.paragraph_format.left_indent = indent
        _set_red(c)

    val = ans
    if val and "$" not in val and re.search(r"\\[a-zA-Z]|\^\{|_\{", val):
        val = f"${val}$"                          # bare latex fragment from tail_extract
    c = _row(t, no)
    p = c.paragraphs[0]
    p.add_run(pre + "Ans: ").bold = True
    write_rich_line(p, sub(answer_display(e, val, mcq_label)) if val else MISSING)
    _close(c)

    c = _row(t, "")
    if sol:
        add_content(c, sub(sol), dirs, [])
        _prefix_bold(c, pre + "Solution: ")
    else:
        c.paragraphs[0].add_run(pre + "Solution: ").bold = True
        c.paragraphs[0].add_run(MISSING)
    c.add_paragraph()                             # gap under the slot, so pairs read apart
    _close(c)


BODY_FONT = "Times New Roman"
PAGE = (Cm(21.0), Cm(29.7))                                                # A4 portrait
MARGINS = {"top": Cm(1), "bottom": Cm(1), "right": Cm(1), "left": Cm(2)}   # left = marking room


def _page_setup(doc) -> None:
    """Print setup for every export: A4, serif body throughout and the worksheet margins.

    The font goes on the styles, not on runs, so everything that inherits them — sub-parts,
    table cells, options, captions — picks it up; `w:rFonts` is set directly because setting
    only `font.name` leaves the east-asian and complex-script slots on the theme font, which
    is what makes stray glyphs render in the wrong face."""
    for name in ("Normal", "Heading 1", "Heading 2"):
        try:
            st = doc.styles[name]
        except KeyError:
            continue
        st.font.name = BODY_FONT
        rpr = st.element.get_or_add_rPr()
        rf = rpr.find(qn("w:rFonts"))
        if rf is None:
            rf = rpr.makeelement(qn("w:rFonts"), {})
            rpr.append(rf)
        for slot in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
            rf.set(qn(slot), BODY_FONT)
    doc.styles["Normal"].font.size = Pt(11)
    for sec in doc.sections:
        sec.page_width, sec.page_height = PAGE
        for side, val in MARGINS.items():
            setattr(sec, f"{side}_margin", val)


def build_docx(ctx: context.Ctx, entries: list[dict], title: str,
               with_solutions: bool, out_path: Path,
               mcq_label: str = "letter_bare", blank: str = "dots",
               marks_col: bool = False, caption: dict | None = None,
               sections: bool = True, show_total: bool = True, teacher: bool = False) -> Path:
    global BLANK_FILL, CAPTION
    BLANK_FILL = "_" * 12 if blank == "underscore" else "." * 16
    CAPTION = {**CAPTION_DEFAULTS, **{k: v for k, v in (caption or {}).items() if v is not None}}
    running: dict = {}                              # continuous caption counters (whole doc)
    doc = Document()
    _page_setup(doc)
    # the grid spans the full text width, so it follows the margins rather than a fixed guess
    sec = doc.sections[0]
    text_w = Emu(sec.page_width - sec.left_margin - sec.right_margin)
    widths = ((Cm(1.2), Emu(text_w - Cm(1.2) - Cm(1.5)), Cm(1.5)) if marks_col
              else (Cm(1.3), Emu(text_w - Cm(1.3))))
    doc.add_heading((title or "Worksheet") + (" — Answers (teacher copy)" if teacher else ""), level=1)
    grand = total_marks(entries)
    if show_total and grand:                         # whole-paper total, under the title
        p = doc.add_paragraph()
        p.add_run(f"Total: {grand} marks").bold = True
    if not teacher:
        doc.add_paragraph("Name: ____________    Class: ______    Date: __________")

    if sections:
        # group into Section A/B/… by type, sorted by marks within; numbering stays
        # continuous across sections. Reorder `entries` to the printed order so the answers
        # section below (same continuous numbers) matches.
        i = 0
        ordered = []
        for label, name, group in order_sections(entries):
            doc.add_heading(f"Section {label} — {name}", level=2)
            t = _grid(doc, widths)
            for e in group:
                i += 1
                _question_block(t, i, e, ctx, mcq_label, marks_col, running, teacher=teacher)
                ordered.append(e)
        entries = ordered
    else:
        t = _grid(doc, widths)
        for i, e in enumerate(entries, 1):
            _question_block(t, i, e, ctx, mcq_label, marks_col, running, teacher=teacher)

    if with_solutions and not teacher:               # end answers table (teacher copy inlines instead)
        doc.add_page_break()
        doc.add_heading("Answers & Solutions", level=1)
        st = _grid(doc, (Cm(1.3), Emu(text_w - Cm(1.3))))
        for i, e in enumerate(entries, 1):
            dirs = entry_img_dirs(ctx, e)
            no = f"{i}."                                 # question number on its first row only
            for label, ans, sol in answer_slots(e):      # same slots, same pair, as the marking copy
                _answer_pair(st, e, i, mcq_label, dirs, ans, sol, label=label, no=no)
                no = ""

    tmp = Path(tempfile.mkdtemp(prefix="questgen_docx_"))
    try:
        tmp_f = tmp / (out_path.name)
        doc.save(tmp_f)
        out_path.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(tmp_f, out_path)
    finally:
        shutil.rmtree(tmp, ignore_errors=True)
    return out_path


if __name__ == "__main__":
    import argparse
    import json
    ap = argparse.ArgumentParser(description=__doc__.split("\n")[0])
    ap.add_argument("--qids", required=True, help="comma-separated qids")
    ap.add_argument("--title", default="Worksheet")
    ap.add_argument("--with-solutions", action="store_true")
    ap.add_argument("--out", default="")
    context.add_ctx_args(ap)
    a = ap.parse_args()
    ctx = context.ctx_from_args(a)
    want = a.qids.split(",")
    pool = {}
    for f in sorted(ctx.interim_dir.glob("*.jsonl")):
        if f.name.endswith(".clean.jsonl") or not (ctx.interim_dir / f"{f.stem}.clean.jsonl").is_file():
            import interim_build as ib
            for r in ib.read_jsonl(f):
                pool[r["qid"]] = r
    entries = [pool[q] for q in want if q in pool]
    out = Path(a.out) if a.out else ctx.outputs_dir / f"{a.title}.docx"
    print(build_docx(ctx, entries, a.title, a.with_solutions, out))
